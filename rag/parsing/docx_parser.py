"""Word 文档解析模块。"""

from __future__ import annotations

import os
from typing import Dict, List

from ..processing.cleaners import clean_plain_text
from ..exceptions import DocumentParseError


def extract_text_from_docx(file_path: str) -> str:
    """从 .docx 提取纯文本。"""
    return "\n\n".join(page["text"] for page in extract_pages_from_docx(file_path))


def extract_pages_from_docx(file_path: str) -> List[Dict]:
    """
    Word 没有稳定的物理页概念，这里按文档整体返回 page=None。
    后续分块会继续附带 doc_name / section 等元数据。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("缺少 python-docx：pip install python-docx") from exc

    try:
        doc = Document(file_path)
        parts: List[str] = []

        for para in doc.paragraphs:
            text = clean_plain_text(para.text)
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [clean_plain_text(cell.text) for cell in row.cells]
                line = " | ".join(x for x in cells if x)
                if line:
                    parts.append(line)

        text = clean_plain_text("\n".join(parts))
        return [{
            "page": None,
            "page_number": None,
            "text": text,
            "method": "docx",
            "extraction_method": "docx",
            "char_count": len(text),
            "text_length": len(text),
            "is_ocr": False,
            "is_ocr_text": False,
            "ocr_backend": "",
            "ocr_confidence": None,
            "is_scanned_page": False,
            "failure_reason": "",
            "quality_stats": {},
        }] if text else []
    except Exception as exc:
        raise DocumentParseError(f"Word 文档解析失败: {file_path}: {exc}") from exc
