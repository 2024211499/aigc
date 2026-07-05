# -*- coding: utf-8 -*-
"""
导出服务接口
支持导出学习计划、习题集、试卷为 PDF / Word / Markdown
"""

import json
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..core.config import settings
from ..core.database import get_db, StudyPlan, ExamPaper, Exercise, ExportFile, gen_id
from ..core.exceptions import NotFoundError

router = APIRouter()
logger = logging.getLogger(__name__)


# ─── 字段中文标签映射 ────────────────────────────────────────────
# 覆盖 LLM 输出中所有可能出现的英文 key
_KEY_LABELS: Dict[str, str] = {
    # 学习计划
    "learning_objectives":    "学习目标",
    "prerequisite_knowledge": "前置知识",
    "core_knowledge_points":  "核心知识点",
    "daily_tasks":            "每日任务安排",
    "review_suggestions":     "复习建议",
    "study_tips":             "学习技巧",
    "summary":                "总结",
    "overview":               "概览",
    # 知识点
    "name":                   "名称",
    "description":            "说明",
    "importance":             "重要程度",
    "difficulty":             "难度",
    "examples":               "例题",
    "formulas":               "公式",
    "notes":                  "注意事项",
    # 每日任务
    "day":                    "第几天",
    "tasks":                  "任务内容",
    "time_minutes":           "预计时长（分钟）",
    "focus":                  "学习重点",
    # 题目通用
    "questions":              "题目",
    "exercises":              "习题",
    "question_text":          "题干",
    "question_type":          "题型",
    "stem":                   "题干",
    "options":                "选项",
    "answer":                 "答案",
    "answers":                "答案",
    "explanation":            "解析",
    "explanations":           "解析",
    "score":                  "分值",
    "knowledge_point":        "考查知识点",
    "knowledge_points":       "知识点",
    "hints":                  "提示",
    "solution":               "解题过程",
    # 选项
    "A": "A", "B": "B", "C": "C", "D": "D",
    # 试卷头部
    "instructions":           "答题说明",
    "exam_instructions":      "答题说明",
    # 试卷结构
    "exam_type":              "考试类型",
    "total_questions":        "题目数量",
    "duration_minutes":       "考试时长（分钟）",
    "total_score":            "总分",
    "passing_score":          "及格线",
    "sections":               "大题",
    "section_title":          "大题名称",
    "section_score":          "大题总分",
    # 通用
    "title":                  "标题",
    "content":                "内容",
    "chapter_name":           "章节",
    "course_name":            "课程",
    "created_at":             "生成时间",
    "remark":                 "备注",
}

def _label(key: str) -> str:
    """将英文 key 转换为中文标签；未知 key 做驼峰→空格处理后原样返回。"""
    if key in _KEY_LABELS:
        return _KEY_LABELS[key]
    # 简单处理：下划线换成空格，首字母大写
    return key.replace("_", " ").capitalize()


# ─── 结构化渲染：将任意嵌套内容转为可读段落列表 ─────────────────

def _flatten_to_lines(obj: Any, indent: int = 0) -> List[str]:
    """
    将任意结构（dict/list/scalar）递归展开为可读文本行。
    不输出原始 JSON，只输出人类可读内容。
    """
    lines: List[str] = []
    prefix = "    " * indent

    if isinstance(obj, dict):
        for k, v in obj.items():
            label = _label(k)
            if isinstance(v, (dict, list)):
                lines.append(f"{prefix}{label}：")
                lines.extend(_flatten_to_lines(v, indent + 1))
            elif v not in (None, "", [], {}):
                lines.append(f"{prefix}{label}：{v}")

    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if isinstance(item, dict):
                # 如果是题目 dict，用序号标头
                if any(k in item for k in ("question_text", "stem", "题干")):
                    lines.append(f"{prefix}第 {i + 1} 题")
                    lines.extend(_flatten_to_lines(item, indent + 1))
                elif "day" in item or "tasks" in item:
                    lines.append(f"{prefix}Day {item.get('day', i + 1)}")
                    lines.extend(_flatten_to_lines(
                        {k: v for k, v in item.items() if k != "day"}, indent + 1
                    ))
                else:
                    lines.extend(_flatten_to_lines(item, indent))
            elif item not in (None, ""):
                lines.append(f"{prefix}• {item}")
    elif obj not in (None, ""):
        for part in str(obj).split("\n"):
            if part.strip():
                lines.append(f"{prefix}{part}")

    return lines


# ─── 单题渲染（选择题 / 主观题 统一格式）────────────────────────

def _render_question_lines(item: dict, idx: int, include_answer: bool = True) -> List[str]:
    """将一道题目 dict 渲染为可读行列表。"""
    lines: List[str] = []
    q_type = item.get("question_type", item.get("题型", ""))
    stem   = item.get("question_text", item.get("stem", item.get("题干", "")))
    score  = item.get("score", item.get("分值", ""))

    header = f"第 {idx} 题"
    if q_type:
        header += f"  [{q_type}]"
    if score:
        header += f"  ({score} 分)"
    lines.append(header)

    if stem:
        lines.append(f"    {stem}")

    options = item.get("options", item.get("选项"))
    if options:
        if isinstance(options, dict):
            for k, v in options.items():
                lines.append(f"    {k}. {v}")
        elif isinstance(options, list):
            for j, opt in enumerate(options):
                letter = chr(ord("A") + j)
                lines.append(f"    {letter}. {opt}")

    if include_answer:
        answer = item.get("answer", item.get("答案", ""))
        if answer:
            lines.append(f"    【答案】{answer}")
        explanation = item.get("explanation", item.get("解析", ""))
        if explanation:
            lines.append(f"    【解析】{explanation}")

    kp = item.get("knowledge_point", item.get("知识点", ""))
    if kp:
        lines.append(f"    【考查知识点】{kp}")

    return lines


# ─── 内容预处理：把 str→dict 的 JSON 字段解包 ───────────────────

def _unpack(obj: Any) -> Any:
    """
    递归解包内容字段：
    - JSON 字符串 → parse 为 dict/list（含 markdown 代码块的字符串会先剥除围栏）
    - {"raw": <str|dict>} 单键兜底结构 → 提取并再次解包 raw 值
    - list 中每个元素同样递归处理
    """
    import re as _re

    if isinstance(obj, str):
        stripped = obj.strip()
        # 兼容旧版 bridge.py 生成的 ```json ... ``` 包裹
        if stripped.startswith("```"):
            stripped = _re.sub(r"```(?:json)?\s*", "", stripped, flags=_re.IGNORECASE).strip()
            stripped = _re.sub(r"```\s*$", "", stripped).strip()
        if stripped.startswith(("{", "[")):
            try:
                return _unpack(json.loads(stripped))
            except Exception:
                pass
        return obj

    if isinstance(obj, dict):
        # 解包 _parse_json_response 的兜底格式 {"raw": ...}
        if list(obj.keys()) == ["raw"]:
            return _unpack(obj["raw"])
        return {k: _unpack(v) for k, v in obj.items()}

    if isinstance(obj, list):
        return [_unpack(i) for i in obj]

    return obj


def _ensure_dict(content: Any) -> dict:
    """确保 _unpack 的结果始终是 dict，防止 content.get() 报错。"""
    if isinstance(content, dict):
        return content
    if isinstance(content, list):
        # 题目列表直接放进 questions 键
        return {"questions": content}
    if content not in (None, ""):
        return {"内容": str(content)}
    return {}


# ─── Word 导出 ───────────────────────────────────────────────────

# 已知顶层 section 的处理顺序
_PLAN_SECTIONS = [
    "learning_objectives",
    "prerequisite_knowledge",
    "core_knowledge_points",
    "daily_tasks",
    "review_suggestions",
    "study_tips",
    "summary",
]
_EXAM_SECTIONS = ["questions", "exercises", "sections"]


def _export_to_docx(title: str, content: dict, file_path: str,
                    is_exam: bool = False, include_answers: bool = True):
    """将结构化内容导出为 Word 文档，无 JSON 输出。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        content = _ensure_dict(_unpack(content))
        doc = Document()

        # 标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _add_text_block(lines: List[str]):
            for line in lines:
                stripped = line.lstrip()
                depth = (len(line) - len(stripped)) // 4
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12 * depth)
                if stripped.startswith("第 ") and ("题" in stripped[:10]):
                    run = p.add_run(stripped)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)
                elif stripped.startswith("【") and stripped.endswith("】") or \
                     stripped.startswith("【"):
                    run = p.add_run(stripped)
                    run.bold = True
                else:
                    p.add_run(stripped)

        def _add_section(key: str):
            val = content.get(key)
            if val in (None, [], {}, ""):
                return
            doc.add_heading(_label(key), level=1)

            # 题目列表特殊处理
            if key in ("questions", "exercises") and isinstance(val, list):
                for i, item in enumerate(val, 1):
                    if isinstance(item, dict):
                        lines = _render_question_lines(item, i, include_answer=include_answers)
                        _add_text_block(lines)
                        doc.add_paragraph()   # 题间空行
                    else:
                        doc.add_paragraph(f"{i}. {item}")
            else:
                lines = _flatten_to_lines(val)
                _add_text_block(lines)

        # 按顺序渲染已知 section
        order = _EXAM_SECTIONS if is_exam else _PLAN_SECTIONS
        rendered = set(order)
        for key in order:
            _add_section(key)

        # 渲染剩余未知字段（不再 JSON dump）
        skip = {"title", "course_name", "chapter_name", "answers", "explanations", "raw"}
        for key, val in content.items():
            if key in rendered or key in skip or val in (None, [], {}, ""):
                continue
            doc.add_heading(_label(key), level=1)
            lines = _flatten_to_lines(val)
            _add_text_block(lines)

        doc.save(file_path)
        return True
    except ImportError:
        logger.warning("python-docx 未安装，跳过 Word 导出")
        return False
    except Exception as e:
        logger.error(f"Word 导出失败: {e}", exc_info=True)
        return False


# ─── PDF 导出 ────────────────────────────────────────────────────

def _register_chinese_font():
    """查找并注册系统中文字体，返回字体名；找不到返回 None。"""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        try:
            registered = pdfmetrics.getRegisteredFontNames()
            for name in ("SimHei", "MicrosoftYaHei", "NotoSansCJK", "WenQuanYi", "STHeiti"):
                if name in registered:
                    return name
        except Exception:
            pass

        candidates = [
            ("C:/Windows/Fonts/simhei.ttf",  "SimHei",         None),
            ("C:/Windows/Fonts/msyh.ttf",    "MicrosoftYaHei", None),
            ("C:/Windows/Fonts/msyh.ttc",    "MicrosoftYaHei", 0),
            ("C:/Windows/Fonts/simsun.ttc",  "SimSun",         0),
            ("C:/Windows/Fonts/STKAITI.TTF", "STKaiti",        None),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK", 0),
            ("/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",      "NotoSansCJK", 0),
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",           "WenQuanYi",   0),
            ("/usr/share/fonts/truetype/arphic/uming.ttc",              "ArphicUming", 0),
            ("/System/Library/Fonts/STHeiti Medium.ttc", "STHeiti",    0),
            ("/Library/Fonts/Arial Unicode MS.ttf",      "ArialUnicode", None),
        ]
        for path, name, idx in candidates:
            if not os.path.exists(path):
                continue
            try:
                if idx is not None:
                    pdfmetrics.registerFont(TTFont(name, path, subfontIndex=idx))
                else:
                    pdfmetrics.registerFont(TTFont(name, path))
                logger.info(f"已注册中文字体: {name} ({path})")
                return name
            except Exception as e:
                logger.debug(f"注册字体 {path} 失败: {e}")
    except Exception as e:
        logger.warning(f"中文字体注册异常: {e}")
    return None


def _export_to_pdf_simple(title: str, content: dict, file_path: str,
                           is_exam: bool = False, include_answers: bool = True):
    """使用 reportlab 导出 PDF，无 JSON 输出。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.lib.enums import TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import ParagraphStyle

        content = _ensure_dict(_unpack(content))
        font_name = _register_chinese_font() or "Helvetica"

        style_title = ParagraphStyle(
            "CTitle", fontName=font_name, fontSize=18,
            leading=24, alignment=TA_CENTER, spaceAfter=14,
        )
        style_h1 = ParagraphStyle(
            "CH1", fontName=font_name, fontSize=13,
            leading=18, spaceBefore=12, spaceAfter=6,
            textColor=colors.HexColor("#1a3c5e"),
        )
        style_q_header = ParagraphStyle(
            "CQHeader", fontName=font_name, fontSize=11,
            leading=16, spaceBefore=8, spaceAfter=4,
            textColor=colors.HexColor("#2c5282"), leftIndent=0,
        )
        style_normal = ParagraphStyle(
            "CNormal", fontName=font_name, fontSize=10,
            leading=15, spaceAfter=2, leftIndent=10,
        )
        style_indent = ParagraphStyle(
            "CIndent", fontName=font_name, fontSize=10,
            leading=15, spaceAfter=2, leftIndent=25,
        )
        style_answer = ParagraphStyle(
            "CAnswer", fontName=font_name, fontSize=10,
            leading=15, spaceAfter=2, leftIndent=25,
            textColor=colors.HexColor("#276749"),
        )
        style_bullet = ParagraphStyle(
            "CBullet", fontName=font_name, fontSize=10,
            leading=15, spaceAfter=2, leftIndent=20,
        )

        def _esc(text: str) -> str:
            return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        story: list = [
            Paragraph(_esc(title), style_title),
            HRFlowable(width="100%", thickness=1, color=colors.grey),
            Spacer(1, 8),
        ]

        def _render_lines_to_story(lines: List[str]):
            for line in lines:
                stripped = line.lstrip()
                depth = (len(line) - len(stripped)) // 4
                if stripped.startswith("第 ") and "题" in stripped[:10]:
                    story.append(Paragraph(_esc(stripped), style_q_header))
                elif stripped.startswith("【答案】") or stripped.startswith("【解析】"):
                    story.append(Paragraph(_esc(stripped), style_answer))
                elif stripped.startswith("•"):
                    story.append(Paragraph(_esc(stripped), style_bullet))
                elif depth >= 1:
                    story.append(Paragraph(_esc(stripped), style_indent))
                else:
                    story.append(Paragraph(_esc(stripped), style_normal))

        def _add_section(key: str):
            val = content.get(key)
            if val in (None, [], {}, ""):
                return
            story.append(Paragraph(_esc(_label(key)), style_h1))

            if key in ("questions", "exercises") and isinstance(val, list):
                for i, item in enumerate(val, 1):
                    if isinstance(item, dict):
                        lines = _render_question_lines(item, i, include_answer=include_answers)
                        _render_lines_to_story(lines)
                        story.append(Spacer(1, 6))
                    else:
                        story.append(Paragraph(_esc(f"{i}. {item}"), style_normal))
            else:
                lines = _flatten_to_lines(val)
                _render_lines_to_story(lines)

        order = _EXAM_SECTIONS if is_exam else _PLAN_SECTIONS
        rendered = set(order)
        for key in order:
            _add_section(key)

        skip = {"title", "course_name", "chapter_name", "answers", "explanations", "raw"}
        for key, val in content.items():
            if key in rendered or key in skip or val in (None, [], {}, ""):
                continue
            story.append(Paragraph(_esc(_label(key)), style_h1))
            lines = _flatten_to_lines(val)
            _render_lines_to_story(lines)

        doc = SimpleDocTemplate(
            file_path, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=20 * mm, bottomMargin=20 * mm,
        )
        doc.build(story)
        return True

    except Exception as e:
        logger.error(f"PDF 导出失败: {e}", exc_info=True)
        try:
            content = _ensure_dict(_unpack(content))
            lines = [title, "=" * 50]
            lines.extend(_flatten_to_lines(content))
            txt_path = file_path.replace(".pdf", "_fallback.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            import shutil
            shutil.copy(txt_path, file_path)
        except Exception:
            pass
        return False


# ─── Markdown 导出 ────────────────────────────────────────────────

def _to_md(obj: Any, level: int = 2,
           include_answers: bool = True,
           _is_question_list: bool = False) -> List[str]:
    """将任意结构递归转换为 Markdown 行，使用中文标签，不输出 JSON。"""
    lines: List[str] = []

    if isinstance(obj, dict):
        for k, v in obj.items():
            label = _label(k)
            if k in ("answers", "explanations") and not include_answers:
                continue
            if isinstance(v, (dict, list)):
                is_q = k in ("questions", "exercises")
                lines.append(f"{'#' * level} {label}")
                lines.append("")
                lines.extend(_to_md(v, level + 1, include_answers, _is_question_list=is_q))
            elif v not in (None, "", [], {}):
                lines.append(f"**{label}**：{v}")
                lines.append("")

    elif isinstance(obj, list):
        if _is_question_list:
            # 题目列表：带序号的结构化渲染
            for i, item in enumerate(obj, 1):
                if isinstance(item, dict):
                    q_lines = _render_question_lines(item, i, include_answer=include_answers)
                    # 第一行是标头，其余缩进
                    for j, line in enumerate(q_lines):
                        stripped = line.lstrip()
                        if j == 0:
                            lines.append(f"#### {stripped}")
                        elif stripped.startswith("【"):
                            lines.append(f"> {stripped}")
                        elif stripped.startswith(("A.", "B.", "C.", "D.")):
                            lines.append(f"  {stripped}")
                        elif stripped.startswith("•"):
                            lines.append(f"- {stripped[1:].strip()}")
                        else:
                            lines.append(f"  {stripped}")
                    lines.append("")
                else:
                    lines.append(f"{i}. {item}")
                    lines.append("")
        else:
            for item in obj:
                if isinstance(item, dict):
                    # 每日任务等结构：展开
                    if "day" in item or "tasks" in item:
                        day = item.get("day", "")
                        lines.append(f"**第 {day} 天**" if day else "")
                        sub = {k: v for k, v in item.items() if k != "day"}
                        lines.extend(_to_md(sub, level, include_answers))
                    else:
                        lines.extend(_to_md(item, level, include_answers))
                elif item not in (None, ""):
                    lines.append(f"- {item}")
            lines.append("")

    elif obj not in (None, ""):
        lines.append(str(obj))
        lines.append("")

    return lines


# ─── 接口 ────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    format: Optional[str] = "pdf"       # pdf / docx / markdown
    version: Optional[str] = "student"  # student / teacher


@router.post("/export/study-plan/{plan_id}", summary="导出学习计划")
def export_study_plan(
    plan_id: str,
    req: ExportRequest,
    db: Session = Depends(get_db),
):
    plan = db.query(StudyPlan).filter(StudyPlan.id == plan_id).first()
    if not plan:
        raise NotFoundError("学习计划")

    title   = plan.title or "学习计划"
    content = _ensure_dict(_unpack(plan.content or {}))
    ext     = "docx" if req.format == "docx" else "pdf"
    file_name = f"study_plan_{plan_id[:8]}.{ext}"
    file_path = str(settings.EXPORT_DIR / file_name)

    if req.format == "docx":
        _export_to_docx(title, content, file_path)
    else:
        _export_to_pdf_simple(title, content, file_path)

    _record_export(db, "study_plan", plan_id, req.format, file_path, file_name)

    if Path(file_path).exists():
        return FileResponse(path=file_path, filename=file_name,
                            media_type="application/octet-stream")
    return {"success": False, "error": "导出文件生成失败"}


@router.post("/export/exam/{exam_id}", summary="导出试卷")
def export_exam(
    exam_id: str,
    req: ExportRequest,
    db: Session = Depends(get_db),
):
    paper = db.query(ExamPaper).filter(ExamPaper.id == exam_id).first()
    if not paper:
        raise NotFoundError("试卷")

    include_answers = (req.version == "teacher")
    content = _ensure_dict(_unpack(paper.content or {}))
    title   = paper.title or "试卷"
    ext     = "docx" if req.format == "docx" else "pdf"
    suffix  = "teacher" if include_answers else "student"
    file_name = f"exam_{exam_id[:8]}_{suffix}.{ext}"
    file_path = str(settings.EXPORT_DIR / file_name)

    if req.format == "docx":
        _export_to_docx(title, content, file_path, is_exam=True,
                        include_answers=include_answers)
    else:
        _export_to_pdf_simple(title, content, file_path, is_exam=True,
                              include_answers=include_answers)

    _record_export(db, "exam", exam_id, req.format, file_path, file_name)

    if Path(file_path).exists():
        return FileResponse(path=file_path, filename=file_name,
                            media_type="application/octet-stream")
    return {"success": False, "error": "导出文件生成失败"}


@router.post("/export/exercises/{course_id}", summary="导出习题集")
def export_exercises(
    course_id: str,
    req: ExportRequest,
    db: Session = Depends(get_db),
):
    exercises = db.query(Exercise).filter(Exercise.course_id == course_id).all()
    if not exercises:
        return {"success": False, "error": "该课程暂无习题"}

    include_answers = (req.version == "teacher")
    title = f"习题集（共 {len(exercises)} 题）"

    # 将数据库行统一为题目 dict 格式
    questions = []
    for ex in exercises:
        item: dict = {
            "question_type": ex.question_type or "",
            "difficulty":    ex.difficulty or "",
            "question_text": ex.stem or "",
        }
        if ex.options:
            item["options"] = ex.options
        if include_answers:
            if ex.answer:
                item["answer"] = ex.answer
            if ex.explanation:
                item["explanation"] = ex.explanation
        questions.append(item)

    content = {"questions": questions}
    ext     = "docx" if req.format == "docx" else "pdf"
    suffix  = "teacher" if include_answers else "student"
    file_name = f"exercises_{course_id[:8]}_{suffix}.{ext}"
    file_path = str(settings.EXPORT_DIR / file_name)

    if req.format == "docx":
        _export_to_docx(title, content, file_path, is_exam=True,
                        include_answers=include_answers)
    else:
        _export_to_pdf_simple(title, content, file_path, is_exam=True,
                              include_answers=include_answers)

    _record_export(db, "exercises", course_id, req.format, file_path, file_name)

    if Path(file_path).exists():
        return FileResponse(path=file_path, filename=file_name,
                            media_type="application/octet-stream")
    return {"success": False, "error": "导出文件生成失败"}


@router.post("/export/markdown/{plan_id}", summary="导出 Markdown")
def export_markdown(
    plan_id: str,
    db: Session = Depends(get_db),
):
    plan = db.query(StudyPlan).filter(StudyPlan.id == plan_id).first()
    if not plan:
        raise NotFoundError("学习计划")

    content = _ensure_dict(_unpack(plan.content or {}))
    md_lines = [f"# {plan.title or '学习计划'}", ""]

    # 按固定顺序渲染
    rendered = set()
    for key in _PLAN_SECTIONS:
        val = content.get(key)
        if val in (None, [], {}, ""):
            continue
        md_lines.append(f"## {_label(key)}")
        md_lines.append("")
        is_q = key in ("questions", "exercises")
        md_lines.extend(_to_md(val, level=3, _is_question_list=is_q))
        rendered.add(key)

    skip = {"title", "course_name", "chapter_name"}
    for key, val in content.items():
        if key in rendered or key in skip or val in (None, [], {}, ""):
            continue
        md_lines.append(f"## {_label(key)}")
        md_lines.append("")
        md_lines.extend(_to_md(val, level=3))
        rendered.add(key)

    md_content = "\n".join(md_lines)
    file_name = f"study_plan_{plan_id[:8]}.md"
    file_path = str(settings.EXPORT_DIR / file_name)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    return FileResponse(path=file_path, filename=file_name, media_type="text/markdown")


# ─── 辅助：记录导出文件 ─────────────────────────────────────────

def _record_export(db: Session, source_type: str, source_id: str,
                   fmt: str, file_path: str, file_name: str):
    ef = ExportFile(
        id=gen_id(),
        source_type=source_type,
        source_id=source_id,
        export_format=fmt,
        file_path=file_path,
        file_name=file_name,
        file_size=Path(file_path).stat().st_size if Path(file_path).exists() else 0,
    )
    db.add(ef)
    db.commit()
